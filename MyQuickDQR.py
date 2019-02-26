## Note: the format and style can be customized.

## Process

## Step 1: read data
## import pandas as pd
## import MyQuickDQR as dqr
## mydata = pd.read_csv()

## Step 2: define category columns
## mydata[category_columns] = mydata[category_columns].astype('category')
## keep in mind that some columns may need to change from numeric to category: ZIP, etc.

## Step 3: generate Data Quality Report
## dqr.QuickDQR(mydata, 'xxx.docx')

## Step 4: If the output reads "Fail to add graph for (variable name)", you need to manually make a plot. Sorry for the inconvenience.

import pandas as pd
import numpy as np
import scipy.stats as sps
import matplotlib.pyplot as plt
import seaborn as sns
import sklearn as skl
from docx import Document
from docx.shared import Inches

# Function to map out the summary table for Categorical Variable. 
def MapCategory(cat):
    # Initiate Empty Summary Table for Categorical Variable
    cat_col = list(cat.columns)

    t ={'colname': cat_col, 
        'n_record': cat_col, 
        'percent' : cat_col, 
        'unique_v': cat_col, 
        'mode': cat_col, 
        'count_mode': cat_col}

    cat_table = pd.DataFrame(t)

    for l in range(len(cat_table)):
        cat_table.iloc[l,1] = cat.iloc[:,l].count()
        cat_table.iloc[l,2] = round(cat_table.iloc[l,1] / len(cat) * 100,2)
        cat_table.iloc[l,3] = len(cat.iloc[:,l].unique())
        m = cat.iloc[:,l].value_counts()
        cat_table.iloc[l,4] = m.index[0]
        cat_table.iloc[l,5] = m.iloc[0]
    return cat_table

def DesCategory(cat_table):
    # Description for Categorical Variable
    cat_description = []
    for i in range(len(cat_table)): 
        name = str(cat_table['colname'][i])
        n = str(int(cat_table['n_record'][i]))
        p = str(round(cat_table['percent'][i],2))
        unique_v = str(cat_table['unique_v'][i])
        mode = str(cat_table['mode'][i])
        count = str(cat_table['count_mode'][i])
        cat_description.append(name+' is a categorical variable. '+name+\
                               ' has '+n+' lines of records, and is '+p+\
                               '% populated. '+name+' has '+unique_v +\
                               ' unique categories. The most common category is '\
                               +mode+ ', which occured '+count+' times out of '\
                               +n+' records. ')
    return cat_description

def GraphCategory(cat):
    sns.set_style("whitegrid")
    # Create Category Graph
    cat_col = list(cat.columns)
    for c in cat_col:
        m = cat[c].value_counts()
        name = c + '.png'
        level = len(m)
        comment = []
        try: 
            if level >= 20: 
                comment.append(c)
                if m.iloc[0] / m.iloc[2] >= 8: # If the scale has too big difference
                    plot = cat[c].value_counts().head(20).plot(kind='bar')
                    plot.set_yscale('log')
                    plt.savefig(name,bbox_inches = 'tight')
                    plt.clf()
                else: 
                    plot = cat[c].value_counts().head(20).plot(kind='bar')
                    plt.savefig(name,bbox_inches = 'tight')
                    plt.clf()
            else: 
                if m.iloc[0] / m.iloc[2] >= 8: # If the scale has too big difference
                    plot = cat[c].value_counts().plot(kind='bar')
                    plot.set_yscale('log')
                    plt.savefig(name,bbox_inches = 'tight')
                    plt.clf()
                else: 
                    plot = cat[c].value_counts().plot(kind='bar')
                    plt.savefig(name,bbox_inches = 'tight')
                    plt.clf()
        except:
            print('Fail to create graph for', c, '. Try manually.')

    # Description for Categorical Variable: comment on graphs
    cat_description = []
    for c in cat_col: 
        if c in comment: 
            cat_description.append('Below is a graph showing the destribution of '+c+': ')
        else:
            cat_description.append('Below is a graph showing the destribution of '+c+': (Showing top 20 categories)')
    
    return cat_description

    
# Initiate Summary Table for Numerical Variable
def MapNumeric(num):
    num_col = list(num.columns)

    t ={'colname': num_col, 
        'n_record': num_col, 
        'percent' : num_col, 
        'unique_v': num_col, 
        'n_zero': num_col,
        'mode': num_col, 
        'count_mode': num_col,
        'min': num_col,
        'max': num_col, 
        'mean': num_col,
        'std': num_col, 
        }

    num_table = pd.DataFrame(t)

    # Fill in the Numerical Variable Summary Table
    for l in range(len(num_table)):
        num_table.iloc[l,1] = num.iloc[:,l].count()
        num_table.iloc[l,2] = round(num_table.iloc[l,1] / len(num) * 100,2)
        num_table.iloc[l,3] = len(num.iloc[:,l].unique())
        num_table.iloc[l,4] = sum(num.iloc[:,l] == 0)
        m = num.iloc[:,l].value_counts()
        num_table.iloc[l,5] = m.index[0]
        num_table.iloc[l,6] = m.iloc[0]
        num_table.iloc[l,7] = num.iloc[:,l].min()
        num_table.iloc[l,8] = round(num.iloc[:,l].max(), 2)
        num_table.iloc[l,9] = round(num.iloc[:,l].mean(), 2)
        num_table.iloc[l,10] = round(num.iloc[:,l].std(), 2)
    return num_table

def DesNumeric(num_table):
    # Description for Numerical Variable
    num_description1 = []
    for i in range(len(num_table)): 
        name = str(num_table['colname'][i])
        n = str(int(num_table['n_record'][i]))
        p = str(round(num_table['percent'][i],2))
        unique_v = str(num_table['unique_v'][i])
        n_zero = str(num_table['n_zero'][i])
        mode = str(num_table['mode'][i])
        count = str(num_table['count_mode'][i])
        min_ = str(int(num_table['min'][i]))
        max_ = str(int(num_table['max'][i]))
        avg = str(round(num_table['mean'][i],2))
        std = str(round(num_table['std'][i],2))
        num_description1.append(name+' is a numeric variable. '+name+' has '\
                                +n+' lines of records, and is '+p+'% populated. '\
                                +name+' has '+unique_v +' unique categories. '\
                                +'The most common value is '+mode+ ', occured '\
                                +count+' times. '+name+' has '+n_zero+\
                                ' zero values out of '+n+' lines of records. '\
                                +'The summary statistics and distribution is as follows: '\
                                +'(excluding null value)')
    return num_description1


def GraphNumeric(num):
# Create Graph for Numerical Variable
    num_col = list(num.columns)
    for c in num_col:
        null_remove = num[pd.isnull(num[c]) == False]
        m = null_remove[c].value_counts()
        mode_count = m.iloc[0]
        next_mode_count = m.iloc[4]
        name = c+'.png'
        try:
            if (mode_count / next_mode_count) >= 5: 
                sns.distplot(null_remove[c],bins = 80, kde=False, rug = False).set_yscale('log')
                plt.savefig(name,bbox_inches = 'tight')
                plt.clf()
            else: 
                sns.distplot(null_remove[c],bins = 80, kde=False, rug = False)
                plt.savefig(name,bbox_inches = 'tight')
                plt.clf()
        except:
            print('Fail to create graph for',c,'. Try manually.')
    



def QuickDQR(mydata, filename):    
    # Divide fields in to Category Variable and Numerical Variable
    cat = mydata.loc[:, mydata.dtypes == 'category']
    num = mydata.loc[:, mydata.dtypes != 'category']

    #Produce Cat Results
    cat_table = MapCategory(cat)
    cat_description1 = DesCategory(cat_table)
    cat_description2 = GraphCategory(cat)

    #Produce Num Results
    num_table = MapNumeric(num)
    num_description1 = DesNumeric(num_table)
    GraphNumeric(num)
    
    # Document Output!!!
    document = Document()
    document.add_heading('Data Quality Report', 0)
    
    # High-level summary
    document.add_heading('High-Level Description of the Data',level = 1)
    document.add_paragraph('This dataset shows the information about (dataset name). '\
                           +'It covers the period from (1/1/2010) to (12/31/2010). '\
                           +'The dataset has '+str(mydata.shape[1])+\
                           ' fields and '+str(mydata.shape[0])+' records.', style = 'Body Text')
    # Summary table of all fields
    document.add_heading('Summary Table of All Fields',level = 1)
    document.add_paragraph('After understanding each field, I re-categorized '\
                           +'those fields into numerical and categorical fields. '\
                           +str(len(num_table))+' field is recognized as numerical field '\
                           +'and the rest of the '+str(len(cat_table))+' fields are categorical fields. '\
                           +'The following are two summary tables for categorical fields '\
                           +'and numerical fields followed by each individual field’s '\
                           +'detailed description respectively.',style = 'Body Text')
    
    # Categorical Variable: 
    document.add_heading('Categorical Variable Summary: ', level = 2)

    # Initiate Summary Table Header
    table = document.add_table(rows = 1, cols = 6, style = 'Light Grid Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Number of Records'
    hdr_cells[2].text = 'Populated %'
    hdr_cells[3].text = 'Unique Value'
    hdr_cells[4].text = 'Most Common Category'
    hdr_cells[5].text = 'Occurance of Common Category'

    # Fill in Summary Table Cell
    cat_col = list(cat.columns)
    for i in range(len(cat_col)): 
        row_cells = table.add_row().cells
        for j in range(6): 
            row_cells[j].text = str(cat_table.iloc[i,j])

    # Individual Field: 
    document.add_heading('Individual Fields: ', level = 3)
    
    for i in range(len(cat_description1)):
        name = cat_col[i]
        document.add_paragraph(name, style = 'List Number')
        document.add_paragraph(cat_description1[i], style = 'Body Text')
    
        table = document.add_table(rows = 1, cols = 6, style = 'Light Grid Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Number of Records'
        hdr_cells[2].text = 'Populated %'
        hdr_cells[3].text = 'Unique Value'
        hdr_cells[4].text = 'Most Common Category'
        hdr_cells[5].text = 'Occurance of Common Category'
        row_cells = table.add_row().cells
        for j in range(6):
            row_cells[j].text = str(cat_table.iloc[i,j])
        
        document.add_paragraph(cat_description2[i], style = 'Body Text')
        try: 
            document.add_picture(name+'.png')
        except: 
            print('Fail to add graph for',name,'. Try manually. ')
    

    # Numeric Variable: 
    document.add_heading('Numeric Variable Summary: ', level = 2)

    # Initiate Summary Table Header
    table = document.add_table(rows = 1, cols = 11, style = 'Light Grid Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Number of Records'
    hdr_cells[2].text = 'Populated %'
    hdr_cells[3].text = 'Unique Value'
    hdr_cells[4].text = 'Number of Zero'
    hdr_cells[5].text = 'Most Common Value'
    hdr_cells[6].text = 'Occurance of Common Value'
    hdr_cells[7].text = 'Min'
    hdr_cells[8].text = 'Max'
    hdr_cells[9].text = 'Average'
    hdr_cells[10].text = 'Standard Deviation'

    # Fill in Summary Table Cell
    num_col = list(num.columns)
    for i in range(len(num_col)): 
        row_cells = table.add_row().cells
        for j in range(11): 
            row_cells[j].text = str(num_table.iloc[i,j])
        
    # Individual Field: 
    document.add_heading('Individual Fields: ', level = 3)

    for i in range(len(num_description1)):
        name = num_col[i]
        document.add_paragraph(name, style = 'List Number')
        document.add_paragraph(num_description1[i], style = 'Body Text')
    
        table = document.add_table(rows = 1, cols = 11, style = 'Light Grid Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Number of Records'
        hdr_cells[2].text = 'Populated %'
        hdr_cells[3].text = 'Unique Value'
        hdr_cells[4].text = 'Number of Zero'
        hdr_cells[5].text = 'Most Common Value'
        hdr_cells[6].text = 'Occurance of Common Value'
        hdr_cells[7].text = 'Min'
        hdr_cells[8].text = 'Max'
        hdr_cells[9].text = 'Average'
        hdr_cells[10].text = 'Standard Deviation'
        row_cells = table.add_row().cells
        for j in range(11):
            row_cells[j].text = str(num_table.iloc[i,j])
    
        try: 
            document.add_picture(name+'.png')
        except: 
            print('Fail to add graph for',name,'. Try manually. ')
        document.save(filename)
    
